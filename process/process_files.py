from docx import Document
import sys
import re
import json
import os
import argparse
import ftfy

def fix(string):
    string = ftfy.fix_text(string)
    string = re.sub(r'[\u200b\u200c\u200d\u2060\ufeff\u8203]', ' ', string)
    
    return string.strip()

def number_to_chinese(num):
    if num >= 10 and num < 20:
        ten_flag = True
    else:
        ten_flag = False

    digit_map = {0: '零', 1: '一', 2: '二', 3: '三', 4: '四', 5: '五', 6: '六', 7: '七', 8: '八', 9: '九'}
    weight_map = ['', '十', '百', '千', '万', '十万', '百万']

    def _get_chinese(digit, weight):
        if digit == 0:
            return digit_map[0]
        elif ten_flag and weight == 1:
            return weight_map[weight]
        else: 
            return digit_map[digit] + weight_map[weight]

    result = ''

    weight = 0
    start = False
    has_follow_zero = False

    while num:
        digit = num % 10

        if digit != 0:
            start = True

        if start:
            if digit != 0:
                result = _get_chinese(digit, weight) + (_get_chinese(0, None) if has_follow_zero else '') + result
                has_follow_zero = False
            else:
                has_follow_zero = True

        num //= 10
        weight += 1
    
    return result

def convert_docx_to_dataset(file_path, dataset_path):
    try:
        doc = Document(file_path)
    except Exception as e:
        sys.stderr.write(f"Failed to open {file_path}\n")
        return 

    for paragraph in doc.paragraphs:
        paragraph.text = fix(paragraph.text) 

    i = 0
    for i in range(len(doc.paragraphs)):
        if doc.paragraphs[i].text:
            title = ''.join(doc.paragraphs[i].text.splitlines())

            if title:
                print(title)
                return
                break
    
    if title.endswith('决定') or title.endswith('解释'):
        sys.stderr.write(f"Skip {file_path}: decision or explanation instead of law\n")
        return

    law_idx = 1
    laws = []
    current_law = None
    for i in range(i + 1, len(doc.paragraphs)):
        pattern = re.compile(f'第{number_to_chinese(law_idx)}条')
        match = pattern.match(doc.paragraphs[i].text)

        # print(pattern, doc.paragraphs[i].text.strip())

        if match:
            if current_law:
                laws.append({
                    'title': title,
                    'article_number': current_law,
                    'content': law_content
                })
            start, end = match.span()
            
            current_law = doc.paragraphs[i].text[start:end]
            law_content = doc.paragraphs[i].text[end:].strip()      
            # print(current_law + '#' + law_content + '#')
            law_idx += 1
        else:
            if doc.paragraphs[i].text and current_law:
                law_content += '\n' + doc.paragraphs[i].text
    
    if current_law:
        laws.append({
            'title': title,
            'article_number': current_law,
            'content': law_content
        })
    else:
        sys.stderr.write(f"no law found in {file_path}, 《{title}》\n")

    with open(dataset_path, 'a', encoding='utf-8') as jsonl_file:
        for law in laws:
            jsonl_file.write(json.dumps(law, ensure_ascii=False) + '\n')



def main(file_dir_path, dataset_path):
    if os.path.exists(dataset_path):
        os.remove(dataset_path)

    for file_name in os.listdir(file_dir_path):
        file_path = os.path.join(file_dir_path, file_name)
        convert_docx_to_dataset(file_path, dataset_path)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--file_dir_path', type=str, required=True)
    parser.add_argument('--dataset_path', type=str, required=True) 
    args = parser.parse_args()

    main(args.file_dir_path, args.dataset_path)